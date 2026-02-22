from pypdf import PdfReader
from docx import Document
from typing import List, Dict, Any
import os
import time
from services.document_types import normalize_extension
from services.ner_service import extract_entities


def _ner_progress(idx: int, total: int, t0: float) -> None:
    """Print NER progress every 10% or every 50 chunks (whichever fires first)."""
    step = max(1, min(50, total // 10))
    if (idx + 1) % step == 0 or idx + 1 == total:
        elapsed = time.time() - t0
        pct = (idx + 1) * 100 // total
        filled = pct // 5
        bar = "█" * filled + "░" * (20 - filled)
        print(f"  [NER] [{bar}] {pct:3d}% ({idx+1}/{total})  {elapsed:.1f}s elapsed", flush=True)


class DocumentIndexer:

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap


    def extract_text_from_pdf(self, pdf_path: str) -> str:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def extract_text_from_docx(self, docx_path: str) -> str:
        doc = Document(docx_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    text_parts.append("\t".join(row_cells))

        return "\n".join(text_parts)

    def chunk_text(self, text: str) -> List[str]:
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)

                if break_point > self.chunk_size // 2:
                    chunk = text[start:start + break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - self.chunk_overlap

        return [c for c in chunks if c]

    async def process_pdf(self, pdf_path: str, document_id: str = None, original_filename: str = None) -> tuple[List[str], List[Dict[str, Any]]]:
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF файл не найден: {pdf_path}")

        text = self.extract_text_from_pdf(pdf_path)
        chunks = self.chunk_text(text)

        if document_id is None:
            import hashlib
            document_id = hashlib.sha256(os.path.basename(pdf_path).encode()).hexdigest()[:16]

        source_name = original_filename if original_filename else os.path.basename(pdf_path)

        print(f"Извлечено {len(chunks)} чанков из {pdf_path}")
        print(f"[NER] Начинаю обработку {len(chunks)} чанков...", flush=True)
        t0 = time.time()
        metadata = []
        for idx, chunk in enumerate(chunks):
            ner = extract_entities(chunk)
            metadata.append({
                "document_id": document_id,
                "source": source_name,
                "chunk_id": idx,
                "total_chunks": len(chunks),
                "entity_texts": ner["entity_texts"],
                "entity_labels": ner["entity_labels"],
            })
            _ner_progress(idx, len(chunks), t0)

        print(f"[NER] Готово за {time.time() - t0:.1f}s", flush=True)
        return chunks, metadata

    async def process_docx(self, docx_path: str, document_id: str = None, original_filename: str = None) -> tuple[List[str], List[Dict[str, Any]]]:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")

        text = self.extract_text_from_docx(docx_path)
        chunks = self.chunk_text(text)

        if document_id is None:
            import hashlib
            document_id = hashlib.sha256(os.path.basename(docx_path).encode()).hexdigest()[:16]

        source_name = original_filename if original_filename else os.path.basename(docx_path)

        print(f"Extracted {len(chunks)} chunks from {docx_path}")
        print(f"[NER] Начинаю обработку {len(chunks)} чанков...", flush=True)
        t0 = time.time()
        metadata = []
        for idx, chunk in enumerate(chunks):
            ner = extract_entities(chunk)
            metadata.append({
                "document_id": document_id,
                "source": source_name,
                "chunk_id": idx,
                "total_chunks": len(chunks),
                "entity_texts": ner["entity_texts"],
                "entity_labels": ner["entity_labels"],
            })
            _ner_progress(idx, len(chunks), t0)

        print(f"[NER] Готово за {time.time() - t0:.1f}s", flush=True)
        return chunks, metadata

    async def process_document(self, file_path: str, document_id: str = None, original_filename: str = None) -> tuple[List[str], List[Dict[str, Any]]]:
        ext = normalize_extension(file_path)
        if ext == ".pdf":
            return await self.process_pdf(file_path, document_id=document_id, original_filename=original_filename)
        if ext == ".docx":
            return await self.process_docx(file_path, document_id=document_id, original_filename=original_filename)

        raise ValueError(f"Unsupported document type: {ext}")

    async def process_multiple_documents(self, file_paths: List[str], document_ids: List[str] = None) -> tuple[List[str], List[Dict[str, Any]]]:
        all_chunks = []
        all_metadata = []

        for idx, file_path in enumerate(file_paths):
            document_id = document_ids[idx] if document_ids and idx < len(document_ids) else None
            chunks, metadata = await self.process_document(file_path, document_id=document_id)
            all_chunks.extend(chunks)
            all_metadata.extend(metadata)

        return all_chunks, all_metadata

    async def process_multiple_pdfs(self, pdf_paths: List[str], document_ids: List[str] = None) -> tuple[List[str], List[Dict[str, Any]]]:
        return await self.process_multiple_documents(pdf_paths, document_ids=document_ids)
